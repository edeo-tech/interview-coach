from fastapi import Request, HTTPException
from typing import Optional, Dict, List, Any
from datetime import datetime, timezone

from crud._generic._db_actions import createDocument, getDocument, getMultipleDocuments, updateDocument
from models.interviews.cv_profile import CVProfile
from services.cv_processing_service import CVProcessingService

async def create_cv_profile(req: Request, user_id: str, file_content: bytes, content_type: str, filename: str) -> CVProfile:
    """Create a new CV profile for a user using OpenAI processing"""
    
    # Initialize CV processing service
    cv_service = CVProcessingService()
    
    try:
        # Process CV using OpenAI
        openai_data = await cv_service.process_cv(file_content, content_type, filename)
        
        # Convert to CVProfile structure
        cv_data = _convert_openai_to_cv_profile(
            user_id=user_id,
            raw_text=_extract_raw_text(file_content, content_type),
            openai_data=openai_data
        )
        
        return await createDocument(req, "cv_profiles", CVProfile, cv_data)
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process CV: {str(e)}"
        )
    finally:
        await cv_service.close()

async def get_user_cv(req: Request, user_id: str) -> Optional[CVProfile]:
    """Get the most recent CV profile for a user"""
    cvs = await getMultipleDocuments(
        req, "cv_profiles", CVProfile, 
        user_id=user_id,
        order_by="parsed_at",
        limit=1
    )
    return cvs[0] if cvs else None

async def update_cv_profile(req: Request, cv_id: str, file_content: bytes, content_type: str, filename: str) -> Optional[CVProfile]:
    """Update an existing CV profile using OpenAI processing"""
    
    # Initialize CV processing service
    cv_service = CVProcessingService()
    
    try:
        # Process CV using OpenAI
        openai_data = await cv_service.process_cv(file_content, content_type, filename)
        
        # Get existing CV to preserve user_id
        existing_cv = await getDocument(req, "cv_profiles", CVProfile, _id=cv_id)
        if not existing_cv:
            raise HTTPException(status_code=404, detail="CV not found")
        
        # Convert to update data
        update_data = _convert_openai_to_update_data(
            raw_text=_extract_raw_text(file_content, content_type),
            openai_data=openai_data
        )
        
        return await updateDocument(req, "cv_profiles", CVProfile, cv_id, **update_data)
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to update CV: {str(e)}"
        )
    finally:
        await cv_service.close()

# Helper functions for OpenAI data conversion

def _extract_raw_text(file_content: bytes, content_type: str) -> str:
    """Extract raw text from file content for storage"""
    import io
    import PyPDF2
    try:
        from docx import Document
    except ImportError:
        try:
            from python_docx import Document
        except ImportError:
            Document = None
    
    try:
        if content_type == 'text/plain':
            return file_content.decode('utf-8')
        elif content_type == 'application/pdf':
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(stream=file_content, filetype="pdf")
                text = ""
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text += page.get_text() + "\n"
                doc.close()
                return text.strip()
            except:
                # Fallback to PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        elif content_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            if Document is None:
                return ""  # Fallback if docx is not available
            doc = Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
    except Exception:
        return ""
    
    return ""

def _convert_openai_to_cv_profile(user_id: str, raw_text: str, openai_data: Dict[str, Any]) -> CVProfile:
    """Convert OpenAI response to CVProfile model"""
    
    # Extract skills data
    skills_data = openai_data.get('skills', {})
    technical_skills = skills_data.get('technical', [])
    programming_languages = skills_data.get('programming_languages', [])
    frameworks = skills_data.get('frameworks', [])
    
    # Combine all technical skills for legacy compatibility
    combined_skills = list(set(technical_skills + programming_languages + frameworks))
    
    # Extract metadata
    metadata = openai_data.get('metadata', {})
    total_years = metadata.get('total_experience_years', 0)
    if isinstance(total_years, str):
        try:
            total_years = int(float(total_years))
        except:
            total_years = 0
    
    return CVProfile(
        user_id=user_id,
        raw_text=raw_text,
        
        # Enhanced structured data
        personal_info=openai_data.get('personal_info', {}),
        professional_summary=openai_data.get('professional_summary', ''),
        
        # Skills breakdown
        technical_skills=technical_skills,
        programming_languages=programming_languages,
        frameworks=frameworks,
        tools=skills_data.get('tools', []),
        soft_skills=skills_data.get('soft_skills', []),
        spoken_languages=skills_data.get('languages', []),
        
        # Experience and education
        experience=openai_data.get('experience', []),
        education=openai_data.get('education', []),
        certifications=openai_data.get('certifications', []),
        projects=openai_data.get('projects', []),
        
        # Additional information
        additional_info=openai_data.get('additional_info', {}),
        
        # Metadata
        total_experience_years=total_years,
        current_level=metadata.get('current_level', 'junior'),
        primary_field=metadata.get('primary_field', ''),
        confidence_score=float(metadata.get('confidence_score', 0.0)),
        
        # Legacy fields for backward compatibility
        skills=combined_skills,
        experience_years=total_years,
        
        # Processing metadata
        parsed_at=datetime.now(timezone.utc),
        processing_method="openai"
    )

def _convert_openai_to_update_data(raw_text: str, openai_data: Dict[str, Any]) -> Dict[str, Any]:
    """Convert OpenAI response to update data dictionary"""
    
    # Extract skills data
    skills_data = openai_data.get('skills', {})
    technical_skills = skills_data.get('technical', [])
    programming_languages = skills_data.get('programming_languages', [])
    frameworks = skills_data.get('frameworks', [])
    
    # Combine all technical skills for legacy compatibility
    combined_skills = list(set(technical_skills + programming_languages + frameworks))
    
    # Extract metadata
    metadata = openai_data.get('metadata', {})
    total_years = metadata.get('total_experience_years', 0)
    if isinstance(total_years, str):
        try:
            total_years = int(float(total_years))
        except:
            total_years = 0
    
    return {
        'raw_text': raw_text,
        
        # Enhanced structured data
        'personal_info': openai_data.get('personal_info', {}),
        'professional_summary': openai_data.get('professional_summary', ''),
        
        # Skills breakdown
        'technical_skills': technical_skills,
        'programming_languages': programming_languages,
        'frameworks': frameworks,
        'tools': skills_data.get('tools', []),
        'soft_skills': skills_data.get('soft_skills', []),
        'spoken_languages': skills_data.get('languages', []),
        
        # Experience and education
        'experience': openai_data.get('experience', []),
        'education': openai_data.get('education', []),
        'certifications': openai_data.get('certifications', []),
        'projects': openai_data.get('projects', []),
        
        # Additional information
        'additional_info': openai_data.get('additional_info', {}),
        
        # Metadata
        'total_experience_years': total_years,
        'current_level': metadata.get('current_level', 'junior'),
        'primary_field': metadata.get('primary_field', ''),
        'confidence_score': float(metadata.get('confidence_score', 0.0)),
        
        # Legacy fields for backward compatibility
        'skills': combined_skills,
        'experience_years': total_years,
        
        # Processing metadata
        'parsed_at': datetime.now(timezone.utc),
        'processing_method': "openai"
    }