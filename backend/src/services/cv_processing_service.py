import json
import base64
import io
import asyncio
from typing import Dict, List, Optional, Any
from decouple import config
import httpx
from fastapi import HTTPException
import PyPDF2
try:
    from docx import Document
except ImportError:
    try:
        from python_docx import Document
    except ImportError:
        Document = None
from PIL import Image
try:
    import fitz  # PyMuPDF for better PDF handling
except ImportError:
    fitz = None
import logging

logger = logging.getLogger(__name__)

# Environment variables
OPENAI_API_KEY = config('OPENAI_API_KEY', default='', cast=str)

class CVProcessingService:
    def __init__(self):
        if not OPENAI_API_KEY:
            raise ValueError("OPENAI_API_KEY environment variable is required")
        
        self.client = httpx.AsyncClient(
            base_url="https://api.openai.com/v1",
            headers={
                "Authorization": f"Bearer {OPENAI_API_KEY}",
                "Content-Type": "application/json"
            },
            timeout=120.0  # Longer timeout for document processing
        )
    
    async def process_cv(self, file_content: bytes, content_type: str, filename: str) -> Dict[str, Any]:
        """
        Process CV file using OpenAI vision/text capabilities with error handling and validation
        Returns structured CV data
        """
        if not file_content:
            raise HTTPException(status_code=400, detail="Empty file content")
        
        # File size validation (max 10MB)
        if len(file_content) > 10 * 1024 * 1024:
            raise HTTPException(
                status_code=400, 
                detail="File too large. Maximum size is 10MB."
            )
        
        try:
            # Convert file to processable format
            if content_type == 'application/pdf':
                text_content, images = await self._extract_pdf_content(file_content)
            elif content_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                text_content = await self._extract_docx_content(file_content)
                images = []
            elif content_type == 'text/plain':
                text_content = file_content.decode('utf-8')
                images = []
            else:
                raise HTTPException(
                    status_code=400,
                    detail="Unsupported file type. Please upload PDF, DOC, DOCX, or TXT files."
                )
            
            # Validate extracted content
            if not text_content.strip() and not images:
                raise HTTPException(
                    status_code=400,
                    detail="No readable content found in the file. Please ensure the file is not corrupted."
                )
            
            # Use OpenAI to extract structured data with retry logic
            try:
                if images and len(images) > 0:
                    # Use vision API for PDFs with complex layouts
                    cv_data = await self._process_with_vision_retry(text_content, images)
                else:
                    # Use text-based processing for simpler formats
                    cv_data = await self._process_with_text_retry(text_content)
                
                # Validate OpenAI response
                self._validate_cv_data(cv_data)
                return cv_data
                
            except httpx.HTTPStatusError as e:
                logger.error(f"OpenAI API error: {e.response.status_code} - {e.response.text}")
                if e.response.status_code == 429:
                    raise HTTPException(
                        status_code=429,
                        detail="CV processing service is temporarily busy. Please try again in a few moments."
                    )
                elif e.response.status_code == 401:
                    raise HTTPException(
                        status_code=500,
                        detail="CV processing service configuration error. Please contact support."
                    )
                else:
                    raise HTTPException(
                        status_code=500,
                        detail="CV processing service error. Please try again or contact support."
                    )
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Unexpected error processing CV {filename}: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Unexpected error processing CV. Please try again or contact support."
            )
    
    async def _extract_pdf_content(self, content: bytes) -> tuple[str, List[str]]:
        """Extract both text and images from PDF"""
        try:
            # Extract text using PyMuPDF (more reliable than PyPDF2)
            doc = fitz.open(stream=content, filetype="pdf")
            text_content = ""
            images = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text_content += page.get_text() + "\n"
                
                # Convert page to image for vision processing
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better quality
                img_data = pix.tobytes("png")
                img_base64 = base64.b64encode(img_data).decode('utf-8')
                images.append(img_base64)
            
            doc.close()
            return text_content.strip(), images
            
        except Exception as e:
            # Fallback to PyPDF2 if PyMuPDF fails
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                return text_content.strip(), []
            except Exception:
                raise Exception(f"Failed to extract PDF content: {str(e)}")
    
    async def _extract_docx_content(self, content: bytes) -> str:
        """Extract text from DOCX file"""
        if Document is None:
            raise Exception("DOCX processing not available. Please install python-docx.")
        
        try:
            doc = Document(io.BytesIO(content))
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + "\t"
                    text_content += "\n"
            
            return text_content.strip()
        except Exception as e:
            raise Exception(f"Failed to extract DOCX content: {str(e)}")
    
    async def _process_with_vision(self, text_content: str, images: List[str]) -> Dict[str, Any]:
        """Process CV using OpenAI Vision API for complex layouts"""
        
        # Use the first page image for vision processing
        first_image = images[0] if images else None
        
        messages = [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": self._get_cv_parsing_prompt()
                    }
                ]
            }
        ]
        
        # Add image if available
        if first_image:
            messages[0]["content"].append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/png;base64,{first_image}",
                    "detail": "high"
                }
            })
        
        # Add text content as well
        if text_content.strip():
            messages[0]["content"].append({
                "type": "text",
                "text": f"\n\nExtracted text content:\n{text_content}"
            })
        
        response = await self.client.post(
            "/chat/completions",
            json={
                "model": "gpt-4o",  # Vision model
                "messages": messages,
                "response_format": {"type": "json_object"},
                "temperature": 0.1,
                "max_tokens": 4000
            }
        )
        
        if response.status_code != 200:
            raise Exception(f"OpenAI API error: {response.status_code} - {response.text}")
        
        result = response.json()
        content = result['choices'][0]['message']['content']
        
        try:
            return json.loads(content)
        except json.JSONDecodeError as e:
            raise Exception(f"Failed to parse OpenAI response as JSON: {str(e)}")
    
    async def _process_with_text(self, text_content: str) -> Dict[str, Any]:
        """Process CV using text-only OpenAI processing"""
        
        messages = [
            {
                "role": "user", 
                "content": f"{self._get_cv_parsing_prompt()}\n\nCV Content:\n{text_content}"
            }
        ]
        
        response = await self.client.post(
            "/chat/completions",
            json={
                "model": "gpt-4o-mini",  # Text model, more cost-effective
                "messages": messages,
                "response_format": {"type": "json_object"},
                "temperature": 0.1,
                "max_tokens": 4000
            }
        )
        
        if response.status_code != 200:
            raise Exception(f"OpenAI API error: {response.status_code} - {response.text}")
        
        result = response.json()
        content = result['choices'][0]['message']['content']
        
        try:
            return json.loads(content)
        except json.JSONDecodeError as e:
            raise Exception(f"Failed to parse OpenAI response as JSON: {str(e)}")
    
    def _get_cv_parsing_prompt(self) -> str:
        """Get the comprehensive CV parsing prompt"""
        return """
You are an expert CV/Resume parser. Analyze the provided CV/Resume document (image and/or text) and extract structured information in JSON format.

Return ONLY valid JSON with the following structure:

{
  "personal_info": {
    "name": "Full name",
    "email": "email@domain.com",
    "phone": "phone number",
    "location": "city, country",
    "linkedin": "LinkedIn URL if present",
    "github": "GitHub URL if present",
    "website": "Personal website if present"
  },
  "professional_summary": "Brief professional summary or objective",
  "skills": {
    "technical": ["list of technical skills"],
    "programming_languages": ["list of programming languages"],
    "frameworks": ["list of frameworks/libraries"],
    "tools": ["list of tools and software"],
    "soft_skills": ["list of soft skills"],
    "languages": ["spoken languages with proficiency if mentioned"]
  },
  "experience": [
    {
      "company": "Company name",
      "position": "Job title",
      "location": "Work location",
      "start_date": "YYYY-MM format or YYYY if only year",
      "end_date": "YYYY-MM format, YYYY, or 'Present'",
      "duration": "calculated duration",
      "responsibilities": ["list of key responsibilities"],
      "achievements": ["list of achievements with metrics if available"]
    }
  ],
  "education": [
    {
      "institution": "School/University name",
      "degree": "Degree type and field",
      "location": "Location if mentioned",
      "start_date": "YYYY-MM or YYYY",
      "end_date": "YYYY-MM or YYYY", 
      "gpa": "GPA if mentioned",
      "honors": "Honors/distinctions if mentioned",
      "relevant_coursework": ["relevant courses if listed"]
    }
  ],
  "certifications": [
    {
      "name": "Certification name",
      "issuing_organization": "Organization",
      "issue_date": "YYYY-MM or YYYY",
      "expiry_date": "YYYY-MM or YYYY if applicable",
      "credential_id": "ID if provided"
    }
  ],
  "projects": [
    {
      "name": "Project name",
      "description": "Brief description",
      "technologies": ["technologies used"],
      "url": "Project URL if available",
      "achievements": ["key achievements or metrics"]
    }
  ],
  "additional_info": {
    "volunteer_work": ["volunteer experiences"],
    "publications": ["publications if any"],
    "awards": ["awards and recognition"],
    "interests": ["personal interests/hobbies"]
  },
  "metadata": {
    "total_experience_years": "calculated years",
    "current_level": "junior/mid/senior based on experience",
    "primary_field": "main field/domain",
    "confidence_score": "0.0-1.0 confidence in parsing quality"
  }
}

Instructions:
1. Extract information accurately from both image (if provided) and text content
2. If information is not available, use empty strings or empty arrays
3. For dates, normalize to YYYY-MM format when possible, YYYY when only year available
4. Calculate experience duration and total years automatically
5. Infer seniority level (junior: 0-3 years, mid: 3-8 years, senior: 8+ years)
6. Group skills logically (separate programming languages from frameworks)
7. Extract metrics and achievements with numbers when available
8. Maintain consistent formatting and structure
9. Provide confidence score based on parsing quality and completeness
10. Return ONLY the JSON object, no additional text or explanations
"""

    async def _process_with_vision_retry(self, text_content: str, images: List[str], max_retries: int = 3) -> Dict[str, Any]:
        """Process CV with vision API with retry logic"""
        for attempt in range(max_retries):
            try:
                return await self._process_with_vision(text_content, images)
            except Exception as e:
                if attempt == max_retries - 1:
                    raise
                logger.warning(f"Vision processing attempt {attempt + 1} failed: {str(e)}")
                await asyncio.sleep(2 ** attempt)  # Exponential backoff
    
    async def _process_with_text_retry(self, text_content: str, max_retries: int = 3) -> Dict[str, Any]:
        """Process CV with text API with retry logic"""
        for attempt in range(max_retries):
            try:
                return await self._process_with_text(text_content)
            except Exception as e:
                if attempt == max_retries - 1:
                    raise
                logger.warning(f"Text processing attempt {attempt + 1} failed: {str(e)}")
                await asyncio.sleep(2 ** attempt)  # Exponential backoff
    
    def _validate_cv_data(self, cv_data: Dict[str, Any]) -> None:
        """Validate the structure of OpenAI response"""
        if not isinstance(cv_data, dict):
            raise ValueError("OpenAI response is not a valid dictionary")
        
        # Check for required top-level keys
        required_keys = ['personal_info', 'skills', 'experience', 'education', 'metadata']
        missing_keys = [key for key in required_keys if key not in cv_data]
        
        if missing_keys:
            logger.warning(f"Missing keys in OpenAI response: {missing_keys}")
            # Add default values for missing keys
            if 'personal_info' not in cv_data:
                cv_data['personal_info'] = {}
            if 'skills' not in cv_data:
                cv_data['skills'] = {}
            if 'experience' not in cv_data:
                cv_data['experience'] = []
            if 'education' not in cv_data:
                cv_data['education'] = []
            if 'metadata' not in cv_data:
                cv_data['metadata'] = {
                    'total_experience_years': 0,
                    'current_level': 'junior',
                    'primary_field': '',
                    'confidence_score': 0.5
                }
        
        # Validate data types
        if not isinstance(cv_data.get('personal_info'), dict):
            cv_data['personal_info'] = {}
        
        if not isinstance(cv_data.get('skills'), dict):
            cv_data['skills'] = {}
        
        if not isinstance(cv_data.get('experience'), list):
            cv_data['experience'] = []
        
        if not isinstance(cv_data.get('education'), list):
            cv_data['education'] = []
        
        if not isinstance(cv_data.get('certifications'), list):
            cv_data['certifications'] = []
        
        if not isinstance(cv_data.get('projects'), list):
            cv_data['projects'] = []
    
    async def close(self):
        """Close the HTTP client"""
        await self.client.aclose()