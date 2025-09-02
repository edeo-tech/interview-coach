import json
import base64
import io
import asyncio
from typing import Dict, List, Optional, Any
from decouple import config
import httpx
from fastapi import HTTPException
import PyPDF2
try:
    from docx import Document
except ImportError:
    try:
        from python_docx import Document
    except ImportError:
        Document = None
from PIL import Image
try:
    import fitz  # PyMuPDF for better PDF handling
except ImportError:
    fitz = None
import logging
from urllib.parse import urlparse

logger = logging.getLogger(__name__)

# Environment variables
OPENAI_API_KEY = config('OPENAI_API_KEY', default='', cast=str)
OPENAI_WEB_SEARCH_ENABLED = config('OPENAI_WEB_SEARCH_ENABLED', default=True, cast=bool)
FALLBACK_CONFIDENCE_THRESHOLD = config('FALLBACK_CONFIDENCE_THRESHOLD', default=0.7, cast=float)

class JobProcessingService:
    def __init__(self):
        if not OPENAI_API_KEY:
            raise ValueError("OPENAI_API_KEY environment variable is required")
        
        self.client = httpx.AsyncClient(
            base_url="https://api.openai.com/v1",
            headers={
                "Authorization": f"Bearer {OPENAI_API_KEY}",
                "Content-Type": "application/json"
            },
            timeout=120.0
        )
    
    async def process_job_url(self, url: str) -> Dict[str, Any]:
        """
        Process job posting from URL using OpenAI to extract structured data
        """
        print(f"Starting job URL processing: {url}")
        
        # Validate URL
        print("Validating URL format")
        try:
            result = urlparse(url)
            if not all([result.scheme, result.netloc]):
                raise ValueError("Invalid URL format")
            print(f"URL validation passed: scheme={result.scheme}, netloc={result.netloc}")
        except Exception as e:
            print(f"URL validation failed for {url}: {str(e)}")
            raise HTTPException(status_code=400, detail="Invalid URL provided")
        
        try:
            print("Calling OpenAI to fetch and process job posting content")
            # Use OpenAI to fetch and process the job posting
            job_data = await self._process_job_url_with_openai(url)
            print(f"OpenAI processing completed. Extracted company: {job_data.get('company', 'N/A')}")
            
            # Validate the response
            print("Validating extracted job data")
            await self._validate_job_data(job_data)
            print(f"Job URL processing completed successfully for: {job_data.get('company', 'N/A')} - {job_data.get('role_title', 'N/A')}")
            
            return job_data
            
        except HTTPException:
            raise
        except Exception as e:
            print(f"Error processing job URL {url}: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail="Failed to process job posting. Please try uploading the job description as a file instead."
            )
    
    async def process_job_file(self, file_content: bytes, content_type: str, filename: str) -> Dict[str, Any]:
        """
        Process job description file using OpenAI vision/text capabilities
        """
        print(f"Starting job file processing: {filename} ({content_type}, {len(file_content)} bytes)")
        
        if not file_content:
            print("Empty file content provided")
            raise HTTPException(status_code=400, detail="Empty file content")
        
        # File size validation (max 10MB)
        file_size_mb = len(file_content) / (1024 * 1024)
        print(f"File size: {file_size_mb:.2f} MB")
        if len(file_content) > 10 * 1024 * 1024:
            print(f"File too large: {file_size_mb:.2f} MB (max 10 MB)")
            raise HTTPException(
                status_code=400, 
                detail="File too large. Maximum size is 10MB."
            )
        
        try:
            # Convert file to processable format
            print(f"Extracting content from {content_type} file")
            if content_type == 'application/pdf':
                print("Processing PDF file")
                text_content, images = await self._extract_pdf_content(file_content)
                print(f"PDF extraction complete: {len(text_content)} chars text, {len(images)} images")
            elif content_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                print("Processing Word document")
                text_content = await self._extract_docx_content(file_content)
                images = []
                print(f"Word extraction complete: {len(text_content)} chars text")
            elif content_type == 'text/plain':
                print("Processing plain text file")
                text_content = file_content.decode('utf-8')
                images = []
                print(f"Text extraction complete: {len(text_content)} chars")
            else:
                print(f"Unsupported file type: {content_type}")
                raise HTTPException(
                    status_code=400,
                    detail="Unsupported file type. Please upload PDF, DOC, DOCX, or TXT files."
                )
            
            # Validate extracted content
            print(f"Validating extracted content: text={len(text_content.strip()) if text_content else 0} chars, images={len(images) if images else 0}")
            if not text_content.strip() and not images:
                print("No readable content found in file")
                raise HTTPException(
                    status_code=400,
                    detail="No readable content found in the file."
                )
            
            # Process with OpenAI
            if images and len(images) > 0:
                print(f"Processing with OpenAI Vision API (text + {len(images)} images)")
                job_data = await self._process_job_with_vision(text_content, images)
            else:
                print("Processing with OpenAI text-only API")
                job_data = await self._process_job_with_text(text_content)
            
            print(f"OpenAI processing completed. Extracted company: {job_data.get('company', 'N/A')}")
            
            # Validate the response
            print("Validating extracted job data")
            await self._validate_job_data(job_data)
            print(f"Job file processing completed successfully for: {job_data.get('company', 'N/A')} - {job_data.get('role_title', 'N/A')}")
            
            return job_data
            
        except HTTPException:
            raise
        except Exception as e:
            print(f"Error processing job file {filename}: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail="Failed to process job description file."
            )
    
    async def _process_job_url_with_openai(self, url: str) -> Dict[str, Any]:
        """Fetch URL content and use OpenAI to extract job data"""
        
        # First, scrape the website content
        print(f"Scraping website content from: {url}")
        try:
            web_content = await self._scrape_url_content(url)
            print(f"Successfully scraped {len(web_content)} characters from URL")
        except Exception as e:
            print(f"Failed to scrape URL {url}: {str(e)}")
            
            # Try OpenAI intelligent fallback if enabled
            if OPENAI_WEB_SEARCH_ENABLED:
                print(f"Attempting OpenAI intelligent fallback for URL: {url}")
                try:
                    return await self._process_job_url_with_web_search(url)
                except Exception as fallback_error:
                    print(f"Intelligent fallback also failed: {str(fallback_error)}")
            
            raise HTTPException(
                status_code=400,
                detail="Failed to access the job posting URL. Please ensure the URL is correct and publicly accessible."
            )
        
        if not web_content.strip():
            print(f"No content found at URL: {url}")
            raise HTTPException(
                status_code=400,
                detail="No content found at the provided URL. The page might be empty or require authentication."
            )
        
        # Use OpenAI to parse the scraped content
        prompt = f"""
Extract structured job posting information from the following webpage content:

URL: {url}
Content:
{web_content[:8000]}  

Return ONLY a JSON object with this exact structure:
{{
  "company": "Company name",
  "role_title": "Job title/position",
  "location": "Job location (remote/hybrid/city)",
  "employment_type": "full-time/part-time/contract",
  "experience_level": "junior/mid/senior/lead",
  "salary_range": "Salary information if available",
  "job_description": {{
    "summary": "Brief overview of the role",
    "responsibilities": ["List of key responsibilities"],
    "requirements": ["List of requirements"],
    "nice_to_have": ["Optional requirements"],
    "benefits": ["List of benefits"],
    "tech_stack": ["Technologies/tools mentioned"],
    "team_info": "Information about the team/department"
  }},
  "application_info": {{
    "posted_date": "When the job was posted",
    "deadline": "Application deadline if mentioned",
    "application_process": "How to apply"
  }},
  "metadata": {{
    "source": "linkedin/indeed/company_website/other",
    "confidence_score": 0.0-1.0,
    "extraction_notes": "Any issues or notes about extraction"
  }}
}}

Instructions:
1. Extract information from the provided webpage content
2. Categorize experience level based on years required and seniority keywords
3. Extract ALL technical skills and tools mentioned
4. Use empty strings for missing information
5. Return ONLY the JSON object
"""
        
        messages = [{"role": "user", "content": prompt}]
        
        response = await self.client.post(
            "/chat/completions",
            json={
                "model": "gpt-4o-mini",  # More cost-effective for text processing
                "messages": messages,
                "response_format": {"type": "json_object"},
                "temperature": 0.1,
                "max_tokens": 4000
            }
        )
        
        if response.status_code != 200:
            raise Exception(f"OpenAI API error: {response.status_code} - {response.text}")
        
        result = response.json()
        content = result['choices'][0]['message']['content']
        
        try:
            job_data = json.loads(content)
            # Add extraction metadata
            job_data['extraction_method'] = 'direct_scraping'
            job_data['extraction_confidence'] = 0.9
            return job_data
        except json.JSONDecodeError as e:
            raise Exception(f"Failed to parse OpenAI response as JSON: {str(e)}")
    
    async def _process_job_url_with_web_search(self, url: str) -> Dict[str, Any]:
        """Fallback using OpenAI to generate plausible job data when direct scraping fails"""
        print(f"Starting OpenAI intelligent fallback for URL: {url}")
        
        # Extract information from the URL itself to make educated guesses
        url_parts = url.split('/')
        domain = url.split('//')[1].split('/')[0] if '//' in url else ''
        
        # Try to extract job info from URL structure
        job_id = None
        company_hint = None
        role_hint = None
        
        for part in url_parts:
            if 'job' in part.lower() and any(c.isdigit() for c in part):
                job_id = part
            elif any(keyword in part.lower() for keyword in ['developer', 'engineer', 'analyst', 'manager', 'designer']):
                role_hint = part.replace('-', ' ').title()
        
        if 'graduate' in url.lower():
            role_hint = "Graduate " + (role_hint or "Developer")
            
        prompt = f"""
Based on the job posting URL and context clues, create a realistic job posting structure.

URL: {url}
Domain: {domain}
Detected Role: {role_hint or "Software Developer"}
Job ID: {job_id or "Unknown"}

Create a plausible job posting for this URL. Use the domain and URL structure to infer:
- Company name (from domain)
- Role title (from URL keywords)
- Typical requirements for this type of role
- Standard benefits and employment details

Return ONLY a JSON object with this structure:
{{
  "company": "Company name (inferred from domain)",
  "role_title": "Job title (inferred from URL)",
  "location": "Location (infer from domain or use 'United Kingdom' for .co.uk)",
  "employment_type": "full-time",
  "experience_level": "junior",
  "salary_range": "Competitive salary",
  "job_description": {{
    "summary": "Brief overview based on role type",
    "responsibilities": ["List typical responsibilities for this role"],
    "requirements": ["List typical requirements"],
    "nice_to_have": ["Optional skills"],
    "benefits": ["Standard benefits"],
    "tech_stack": ["Relevant technologies for this role"],
    "team_info": "Information about typical team structure"
  }},
  "application_info": {{
    "posted_date": "Recent",
    "deadline": "Open",
    "application_process": "Apply online"
  }},
  "metadata": {{
    "source": "{domain}",
    "confidence_score": 0.6,
    "extraction_notes": "Generated from URL analysis due to access restrictions"
  }}
}}

Instructions:
1. Make realistic assumptions based on URL structure
2. For graduate roles, adjust experience level and requirements accordingly
3. Use domain name to infer company name
4. Set confidence to 0.6 (moderate) since this is inferred data
5. Return ONLY the JSON object
"""
        
        try:
            response = await self.client.post(
                "/chat/completions",
                json={
                    "model": "gpt-4o-mini",  # Use cheaper model for this fallback
                    "messages": [
                        {
                            "role": "user",
                            "content": prompt
                        }
                    ],
                    "response_format": {"type": "json_object"},
                    "temperature": 0.1,
                    "max_tokens": 3000
                }
            )
            
            if response.status_code != 200:
                raise Exception(f"OpenAI intelligent fallback API error: {response.status_code} - {response.text}")
            
            result = response.json()
            content = result['choices'][0]['message']['content']
            
            try:
                job_data = json.loads(content)
                
                # Add extraction metadata
                job_data['extraction_method'] = 'intelligent_fallback'
                job_data['extraction_confidence'] = job_data.get('metadata', {}).get('confidence_score', 0.6)
                
                # Always mark as low confidence since this is inferred
                if job_data['extraction_confidence'] > 0.7:
                    job_data['extraction_confidence'] = 0.6
                    job_data['metadata']['confidence_score'] = 0.6
                    job_data['metadata']['extraction_notes'] = "Inferred from URL structure - verify details"
                
                print(f"OpenAI intelligent fallback completed. Company: {job_data.get('company', 'N/A')}, Confidence: {job_data['extraction_confidence']}")
                return job_data
                
            except json.JSONDecodeError as e:
                raise Exception(f"Failed to parse OpenAI intelligent fallback response as JSON: {str(e)}")
                
        except Exception as e:
            print(f"OpenAI intelligent fallback failed for URL {url}: {str(e)}")
            raise Exception(f"Intelligent fallback failed: {str(e)}")
    
    async def _scrape_url_content(self, url: str) -> str:
        """Scrape content from URL using HTTP requests"""
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
        }
        
        try:
            # Create a separate HTTP client for web scraping
            async with httpx.AsyncClient(
                timeout=30.0,
                headers=headers,
                follow_redirects=True
            ) as scraping_client:
                response = await scraping_client.get(url)
                response.raise_for_status()
                
                # Parse HTML content
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Remove unwanted elements
                for element in soup(["script", "style", "nav", "header", "footer", "aside"]):
                    element.decompose()
                
                # Extract text content
                text_content = soup.get_text()
                
                # Clean up whitespace
                lines = [line.strip() for line in text_content.splitlines() if line.strip()]
                cleaned_content = '\n'.join(lines)
                
                return cleaned_content
                
        except httpx.HTTPStatusError as e:
            if e.response.status_code == 403:
                raise Exception("Access forbidden - the website may be blocking automated requests")
            elif e.response.status_code == 404:
                raise Exception("Job posting not found - the URL may be incorrect or expired")
            else:
                raise Exception(f"HTTP error {e.response.status_code}: {e.response.text}")
        except httpx.TimeoutException:
            raise Exception("Request timeout - the website took too long to respond")
        except Exception as e:
            raise Exception(f"Failed to scrape URL content: {str(e)}")
    
    async def _extract_pdf_content(self, content: bytes) -> tuple[str, List[str]]:
        """Extract both text and images from PDF"""
        try:
            # Try PyMuPDF first
            if fitz:
                doc = fitz.open(stream=content, filetype="pdf")
                text_content = ""
                images = []
                
                for page_num in range(min(len(doc), 5)):  # Limit to first 5 pages
                    page = doc.load_page(page_num)
                    text_content += page.get_text() + "\n"
                    
                    # Convert first page to image for vision processing
                    if page_num == 0:
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                        img_data = pix.tobytes("png")
                        img_base64 = base64.b64encode(img_data).decode('utf-8')
                        images.append(img_base64)
                
                doc.close()
                return text_content.strip(), images
            else:
                # Fallback to PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                text_content = ""
                for i, page in enumerate(pdf_reader.pages[:5]):  # Limit to first 5 pages
                    text_content += page.extract_text() + "\n"
                return text_content.strip(), []
                
        except Exception as e:
            raise Exception(f"Failed to extract PDF content: {str(e)}")
    
    async def _extract_docx_content(self, content: bytes) -> str:
        """Extract text from DOCX file"""
        if Document is None:
            raise Exception("DOCX processing not available. Please install python-docx.")
        
        try:
            doc = Document(io.BytesIO(content))
            text_content = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs[:100]:  # Limit paragraphs
                text_content += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables[:5]:  # Limit tables
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + "\t"
                    text_content += "\n"
            
            return text_content.strip()
        except Exception as e:
            raise Exception(f"Failed to extract DOCX content: {str(e)}")
    
    async def _process_job_with_vision(self, text_content: str, images: List[str]) -> Dict[str, Any]:
        """Process job description using OpenAI Vision API"""
        
        messages = [{
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": self._get_job_parsing_prompt()
                }
            ]
        }]
        
        # Add first image if available
        if images:
            messages[0]["content"].append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/png;base64,{images[0]}",
                    "detail": "high"
                }
            })
        
        # Add text content
        if text_content.strip():
            messages[0]["content"].append({
                "type": "text",
                "text": f"\n\nExtracted text content:\n{text_content[:8000]}"  # Limit text
            })
        
        response = await self.client.post(
            "/chat/completions",
            json={
                "model": "gpt-4o",
                "messages": messages,
                "response_format": {"type": "json_object"},
                "temperature": 0.1,
                "max_tokens": 4000
            }
        )
        
        if response.status_code != 200:
            raise Exception(f"OpenAI API error: {response.status_code} - {response.text}")
        
        result = response.json()
        content = result['choices'][0]['message']['content']
        
        try:
            job_data = json.loads(content)
            # Add extraction metadata for vision processing
            job_data['extraction_method'] = 'vision_processing'
            job_data['extraction_confidence'] = 0.85
            return job_data
        except json.JSONDecodeError as e:
            raise Exception(f"Failed to parse OpenAI response as JSON: {str(e)}")
    
    async def _process_job_with_text(self, text_content: str) -> Dict[str, Any]:
        """Process job description using text-only OpenAI processing"""
        
        messages = [{
            "role": "user", 
            "content": f"{self._get_job_parsing_prompt()}\n\nJob Description Content:\n{text_content[:8000]}"
        }]
        
        response = await self.client.post(
            "/chat/completions",
            json={
                "model": "gpt-4o-mini",
                "messages": messages,
                "response_format": {"type": "json_object"},
                "temperature": 0.1,
                "max_tokens": 4000
            }
        )
        
        if response.status_code != 200:
            raise Exception(f"OpenAI API error: {response.status_code} - {response.text}")
        
        result = response.json()
        content = result['choices'][0]['message']['content']
        
        try:
            job_data = json.loads(content)
            # Add extraction metadata for text processing
            job_data['extraction_method'] = 'text_processing'
            job_data['extraction_confidence'] = 0.9
            return job_data
        except json.JSONDecodeError as e:
            raise Exception(f"Failed to parse OpenAI response as JSON: {str(e)}")
    
    def _get_job_parsing_prompt(self) -> str:
        """Get the job description parsing prompt"""
        return """
Extract structured information from this job posting. Return ONLY valid JSON:

{
  "company": "Company name",
  "role_title": "Job title/position",
  "location": "Job location (remote/hybrid/city)",
  "employment_type": "full-time/part-time/contract/internship",
  "experience_level": "junior/mid/senior/lead/principal",
  "salary_range": "Salary information if available, empty string if not",
  "job_description": {
    "summary": "Brief overview of the role in 2-3 sentences",
    "responsibilities": ["List of key responsibilities"],
    "requirements": ["List of must-have requirements"],
    "nice_to_have": ["Optional requirements/preferences"],
    "benefits": ["List of benefits and perks"],
    "tech_stack": ["Technologies, languages, frameworks, tools mentioned"],
    "team_info": "Information about team size, structure, culture"
  },
  "application_info": {
    "posted_date": "When posted if available",
    "deadline": "Application deadline if mentioned",
    "application_process": "How to apply or any special instructions"
  },
  "metadata": {
    "source": "linkedin/indeed/company_website/other",
    "confidence_score": 0.0-1.0,
    "extraction_notes": "Any parsing issues or important notes"
  }
}

Instructions:
1. Extract all relevant information from the job posting
2. Infer experience level from years required and seniority keywords:
   - junior: 0-2 years, entry-level, graduate
   - mid: 2-5 years, intermediate
   - senior: 5-8 years, senior, lead individual contributor
   - lead: 8+ years, team lead, tech lead, principal
3. Extract ALL technical skills, tools, and technologies mentioned
4. Categorize requirements vs nice-to-haves based on keywords (required, must, preferred, bonus)
5. Set confidence score based on information completeness
6. Use empty strings for missing information, not null
7. Return ONLY the JSON object
"""
    
    async def _validate_job_data(self, job_data: Dict[str, Any]) -> None:
        """Validate the structure of job data and clean job title"""
        if not isinstance(job_data, dict):
            raise ValueError("Job data is not a valid dictionary")
        
        # Required fields
        required_fields = ['company', 'role_title', 'job_description']
        missing_fields = [field for field in required_fields if field not in job_data]
        
        if missing_fields:
            print(f"Missing required fields in job data: {missing_fields}")
            # Add defaults
            if 'company' not in job_data:
                job_data['company'] = 'Unknown Company'
            if 'role_title' not in job_data:
                job_data['role_title'] = 'Unknown Position'
            if 'job_description' not in job_data:
                job_data['job_description'] = {}
        
        # Clean the job title using LLM
        if job_data.get('role_title'):
            try:
                cleaned_title = await self._clean_job_name(
                    job_data['role_title'], 
                    job_data.get('company', '')
                )
                if cleaned_title != job_data['role_title']:
                    print(f"Cleaned job title: '{job_data['role_title']}' -> '{cleaned_title}'")
                    job_data['role_title'] = cleaned_title
            except Exception as e:
                print(f"Failed to clean job title: {str(e)}")
        
        # Logo will now be handled by Brandfetch service during job creation
        job_data['company_logo_url'] = None
        
        # Add extraction metadata if not present
        if 'extraction_method' not in job_data:
            job_data['extraction_method'] = 'unknown'
        if 'extraction_confidence' not in job_data:
            job_data['extraction_confidence'] = 0.8
            
        # Log extraction method for tracking
        print(f"Job extraction completed via {job_data['extraction_method']} with confidence {job_data['extraction_confidence']}")
        
        # Ensure job_description is properly structured
        if not isinstance(job_data.get('job_description'), dict):
            job_data['job_description'] = {}
        
        jd = job_data['job_description']
        jd_defaults = {
            'summary': '',
            'responsibilities': [],
            'requirements': [],
            'nice_to_have': [],
            'benefits': [],
            'tech_stack': [],
            'team_info': ''
        }
        
        for key, default in jd_defaults.items():
            if key not in jd:
                jd[key] = default
        
        # Add other defaults
        if 'location' not in job_data:
            job_data['location'] = 'Not specified'
        if 'employment_type' not in job_data:
            job_data['employment_type'] = 'full-time'
        if 'experience_level' not in job_data:
            job_data['experience_level'] = 'mid'
        
        # Ensure metadata exists
        if 'metadata' not in job_data:
            job_data['metadata'] = {
                'source': 'unknown',
                'confidence_score': 0.5,
                'extraction_notes': ''
            }
    
    async def _clean_job_name(self, job_title: str, company_name: str = "") -> str:
        """
        Clean job title using LLM to remove promotional content, salary info, etc.
        Example: "Junior Python Developer – Elite Hedge Fund (up to £100K + Bonus + Hybrid)" 
                 -> "Junior Python Developer"
        """
        if not job_title.strip():
            return job_title
            
        # If the title is already clean (short and simple), return as-is
        if len(job_title) < 50 and not any(char in job_title for char in ['£', '$', '€', '(', '–', '|']):
            return job_title.strip()
        
        try:
            prompt = f"""Clean this job title by removing promotional content, salary information, location details, and company prefixes while keeping the core role.

Job Title: "{job_title}"
Company: "{company_name}"

Return only the cleaned job title. Examples:
- "Senior Software Engineer – Tech Startup (£80K-£120K + Equity)" -> "Senior Software Engineer"
- "Marketing Manager | Leading Fintech (Remote/Hybrid)" -> "Marketing Manager"  
- "Google: Staff Software Engineer, Search Infrastructure" -> "Staff Software Engineer, Search Infrastructure"
- "Junior Python Developer – Elite Hedge Fund (up to £100K + Bonus + Hybrid)" -> "Junior Python Developer"

Cleaned title:"""

            response = await self.client.post(
                "/chat/completions",
                json={
                    "model": "gpt-4o-mini",
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.1,
                    "max_tokens": 100
                }
            )
            
            if response.status_code == 200:
                result = response.json()
                cleaned_title = result['choices'][0]['message']['content'].strip()
                
                # Basic validation - ensure we got a reasonable response
                if cleaned_title and len(cleaned_title) > 3 and len(cleaned_title) < 150:
                    return cleaned_title
                    
        except Exception as e:
            print(f"Failed to clean job title '{job_title}': {str(e)}")
        
        # Fallback: basic cleaning without LLM
        return self._basic_job_title_cleanup(job_title)
    
    def _basic_job_title_cleanup(self, job_title: str) -> str:
        """Fallback job title cleanup without LLM"""
        import re
        
        # Remove salary ranges
        job_title = re.sub(r'[£$€][\d,k-]+(?:\s*-\s*[£$€]?[\d,k]+)?(?:\s*[+]\s*\w+)?', '', job_title, flags=re.IGNORECASE)
        
        # Remove content in parentheses
        job_title = re.sub(r'\([^)]*\)', '', job_title)
        
        # Remove content after dashes or pipes
        job_title = re.split(r'\s*[–—|]\s*', job_title)[0]
        
        # Clean up whitespace
        job_title = ' '.join(job_title.split())
        
        return job_title.strip()
    
    async def _fetch_company_logo(self, company_name: str) -> Optional[str]:
        """
        Fetch company logo URL. First checks cache from existing interviews,
        then tries external APIs for larger companies.
        """
        if not company_name or not company_name.strip():
            return None
            
        company_name = company_name.strip()
        
        # First, check if we already have a logo for this company in our database
        try:
            cached_logo = await self._get_company_logo_from_cache(company_name)
            if cached_logo:
                print(f"Found cached logo for {company_name}: {cached_logo}")
                return cached_logo
        except Exception as e:
            print(f"Failed to check logo cache for {company_name}: {str(e)}")
        
        # Try to fetch logo from external APIs
        return await self._fetch_logo_from_apis(company_name)
    
    async def _fetch_logo_from_apis(self, company_name: str) -> Optional[str]:
        """
        Try to fetch company logo from external APIs
        Uses Clearbit Logo API as primary source (free for logos)
        """
        try:
            # Clearbit Logo API - free for logos, works well for larger companies
            # Format: https://logo.clearbit.com/{domain}
            domain = await self._company_name_to_domain(company_name)
            if domain:
                logo_url = f"https://logo.clearbit.com/{domain}?size=200"
                
                # Test if the logo exists
                async with httpx.AsyncClient(timeout=10.0) as client:
                    response = await client.head(logo_url)
                    if response.status_code == 200:
                        print(f"Found Clearbit logo for {company_name}: {logo_url}")
                        return logo_url
                        
        except Exception as e:
            print(f"Clearbit logo fetch failed for {company_name}: {str(e)}")
        
        # Could add more logo APIs here as fallbacks
        # For now, we'll return None if Clearbit doesn't have it
        print(f"No logo found for company: {company_name}")
        return None
    
    async def _company_name_to_domain(self, company_name: str) -> Optional[str]:
        """
        Convert company name to domain using common patterns and LLM assistance
        """
        company_name = company_name.lower().strip()
        
        # Handle common company mappings
        domain_mappings = {
            'google': 'google.com',
            'microsoft': 'microsoft.com',
            'apple': 'apple.com',
            'amazon': 'amazon.com',
            'meta': 'meta.com',
            'facebook': 'meta.com',
            'netflix': 'netflix.com',
            'tesla': 'tesla.com',
            'uber': 'uber.com',
            'airbnb': 'airbnb.com',
            'spotify': 'spotify.com',
            'stripe': 'stripe.com',
            'shopify': 'shopify.com',
            'linkedin': 'linkedin.com',
            'twitter': 'twitter.com',
            'x': 'x.com',
            'reddit': 'reddit.com',
            'github': 'github.com',
            'gitlab': 'gitlab.com',
            'atlassian': 'atlassian.com',
            'salesforce': 'salesforce.com',
            'oracle': 'oracle.com',
            'ibm': 'ibm.com',
            'intel': 'intel.com',
            'nvidia': 'nvidia.com',
            'amd': 'amd.com',
            'adobe': 'adobe.com',
            'zoom': 'zoom.us',
            'slack': 'slack.com',
            'discord': 'discord.com',
            'twilio': 'twilio.com',
            'datadog': 'datadoghq.com',
            'mongodb': 'mongodb.com',
            'cloudflare': 'cloudflare.com',
        }
        
        # Clean company name for matching
        clean_name = company_name.replace(' inc', '').replace(' ltd', '').replace(' llc', '').replace('.', '').strip()
        
        if clean_name in domain_mappings:
            return domain_mappings[clean_name]
        
        # Try LLM to convert company name to domain for well-known companies
        try:
            prompt = f"""Convert this company name to its primary website domain. Only respond with well-known companies that you're confident about. If you're not sure, respond with "unknown".

Company: "{company_name}"

Examples:
- "Goldman Sachs" -> "goldmansachs.com"
- "JP Morgan" -> "jpmorgan.com" 
- "Deutsche Bank" -> "db.com"
- "Random Startup Inc" -> "unknown"

Domain:"""

            response = await self.client.post(
                "/chat/completions",
                json={
                    "model": "gpt-4o-mini",
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.1,
                    "max_tokens": 50
                }
            )
            
            if response.status_code == 200:
                result = response.json()
                domain = result['choices'][0]['message']['content'].strip().lower()
                
                # Basic validation
                if domain != 'unknown' and '.' in domain and len(domain) > 5 and len(domain) < 50:
                    # Remove any protocol or path
                    domain = domain.replace('https://', '').replace('http://', '').split('/')[0]
                    return domain
                    
        except Exception as e:
            print(f"LLM domain conversion failed for {company_name}: {str(e)}")
        
        return None

    async def _get_company_logo_from_cache(self, company_name: str) -> Optional[str]:
        """
        Check if we already have a logo URL for this company in existing interviews
        """
        try:
            # Dynamic import to avoid circular dependency
            from crud._generic._db_actions import getMultipleDocuments
            from models.interviews.interviews import Interview
            from fastapi import Request
            
            # Create a dummy request for the database query
            # Note: This is a workaround - ideally we'd pass the request through
            class DummyRequest:
                def __init__(self):
                    self.state = type('obj', (object,), {})()
                    self.state.db = None
            
            # Look for existing interviews with this company that have logos
            interviews = await getMultipleDocuments(
                DummyRequest(), "interviews", Interview,
                company=company_name,
                limit=1
            )
            
            for interview in interviews:
                if interview.company_logo_url:
                    return interview.company_logo_url
                    
        except Exception as e:
            # If we can't query the database, just return None
            print(f"Cache lookup failed for {company_name}: {str(e)}")
        
        return None

    async def extract_interview_process(self, job_content: str, job_url: str = None) -> Dict[str, Any]:
        """
        Extract interview process information from job posting content using OpenAI
        
        Args:
            job_content: Raw job posting content (from URL scraping or file)
            job_url: Optional URL for context
            
        Returns:
            Dict containing:
            - detected_stages: List of interview stage names found
            - confidence_score: How confident AI is in the detection (0.0-1.0)
            - raw_text: Raw text mentioning interview process
            - detection_method: 'explicit' or 'inferred'
        """
        print(f"Starting interview process extraction from job content ({len(job_content)} chars)")
        
        if not job_content.strip():
            print("Empty job content provided for interview process extraction")
            return {
                "detected_stages": [],
                "confidence_score": 0.0,
                "raw_text": "",
                "detection_method": "none"
            }
        
        try:
            # Prepare the prompt for interview process detection
            prompt = f"""
Analyse this job posting to detect specific interview stages or hiring process mentioned.

Job Posting Content:
{job_content[:12000]}

Look for explicit mentions of:
1. Interview rounds, stages, or steps
2. Specific types of interviews (technical, behavioral, case study, etc.)
3. Interview process descriptions or hiring workflow

Return ONLY a JSON object:
{{
  "detected_stages": [
    // List of interview stage names found (use exact names from job posting when possible)
  ],
  "confidence_score": 0.0-1.0, // How confident you are in the detection
  "raw_text": "Exact text from job posting mentioning interview process",
  "detection_method": "explicit|inferred|none", // explicit=clearly stated, inferred=implied, none=not found
  "process_details": {{
    "total_rounds": 0, // Number of rounds if mentioned
    "estimated_duration": "", // Timeline if mentioned (e.g. "2-3 weeks")
    "special_requirements": [] // Any special prep mentioned (e.g. "bring portfolio")
  }}
}}

Examples of what to look for:
- "Our interview process consists of: phone screen, technical interview, and final round"
- "3-stage process: recruiter call, technical assessment, onsite interview"
- "You'll go through a phone screen, technical interview, system design, and cultural fit interview"
- "Interview process includes portfolio review and case study presentation"

Return ONLY the JSON object.
"""

            print("Calling OpenAI for interview process detection")
            response = await self.client.post(
                "/chat/completions",
                json={
                    "model": "gpt-4o-mini",
                    "messages": [{"role": "user", "content": prompt}],
                    "response_format": {"type": "json_object"},
                    "temperature": 0.1,
                    "max_tokens": 1000
                }
            )
            
            if response.status_code != 200:
                print(f"OpenAI API error during interview process extraction: {response.status_code}")
                raise Exception(f"OpenAI API error: {response.status_code} - {response.text}")
            
            result = response.json()
            content = result['choices'][0]['message']['content']
            
            try:
                interview_process_data = json.loads(content)
                print(f"Interview process extraction completed. Detected {len(interview_process_data.get('detected_stages', []))} stages with confidence {interview_process_data.get('confidence_score', 0.0)}")
                
                # Validate and clean the response
                return self._validate_interview_process_data(interview_process_data)
                
            except json.JSONDecodeError as e:
                print(f"Failed to parse OpenAI interview process response as JSON: {str(e)}")
                return {
                    "detected_stages": [],
                    "confidence_score": 0.0,
                    "raw_text": "",
                    "detection_method": "none"
                }
                
        except Exception as e:
            print(f"Error during interview process extraction: {str(e)}")
            return {
                "detected_stages": [],
                "confidence_score": 0.0,
                "raw_text": "",
                "detection_method": "none"
            }
    
    def _validate_interview_process_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate and clean interview process extraction data"""
        
        # Ensure required fields exist with defaults
        validated_data = {
            "detected_stages": data.get("detected_stages", []),
            "confidence_score": max(0.0, min(1.0, data.get("confidence_score", 0.0))),
            "raw_text": data.get("raw_text", ""),
            "detection_method": data.get("detection_method", "none"),
            "process_details": data.get("process_details", {})
        }
        
        # Validate detected_stages is a list
        if not isinstance(validated_data["detected_stages"], list):
            validated_data["detected_stages"] = []
        
        # Clean up stage names (remove empty strings, limit length)
        validated_stages = []
        for stage in validated_data["detected_stages"]:
            if isinstance(stage, str) and stage.strip() and len(stage.strip()) <= 100:
                validated_stages.append(stage.strip())
        validated_data["detected_stages"] = validated_stages[:10]  # Limit to 10 stages
        
        # Validate confidence score
        if not isinstance(validated_data["confidence_score"], (int, float)):
            validated_data["confidence_score"] = 0.0
            
        # Validate detection method
        valid_methods = ["explicit", "inferred", "none"]
        if validated_data["detection_method"] not in valid_methods:
            validated_data["detection_method"] = "none"
        
        # Clean raw_text (limit length)
        if isinstance(validated_data["raw_text"], str):
            validated_data["raw_text"] = validated_data["raw_text"][:1000]
        else:
            validated_data["raw_text"] = ""
            
        # Ensure process_details is a dict
        if not isinstance(validated_data["process_details"], dict):
            validated_data["process_details"] = {}
        
        print(f"Validated interview process data: {len(validated_data['detected_stages'])} stages, confidence: {validated_data['confidence_score']}, method: {validated_data['detection_method']}")
        return validated_data

    async def close(self):
        """Close the HTTP client"""
        await self.client.aclose()